from __future__ import annotations

import html
import math
import os
import time
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DIAGRAM_DIR = ROOT / "diagram"
DOCS_DIR = ROOT / "docs"
AUDIT_DIR = ROOT / "audit"
SCREENSHOT_DIR = ROOT / "screenshots"
MOCKUP_DIR = SCREENSHOT_DIR / "mockup_lite"
PREVIEW_DIR = DIAGRAM_DIR / "preview_png"


SYSTEM = {
    "domain": "Sistem Informasi Penjualan dan Kredit Emas Berbasis Web pada Toko MahenGold",
    "short": "MahenGold",
    "framework": "CodeIgniter 4, PHP 8, MariaDB/MySQL, Bootstrap 5",
    "roles": ["admin", "pelanggan/calon_nasabah tanpa login"],
    "level0_reason": "Enam proses induk digunakan karena struktur route, controller, service, dan view membentuk enam kelompok fungsi utama: layanan publik katalog-simulasi, autentikasi admin, pengelolaan master, transaksi kredit dan jadwal, pembayaran-piutang, serta laporan-komunikasi WhatsApp.",
}

TABLES = [
    ("users", "Menyimpan akun admin yang dapat masuk ke panel administrasi.", ["id PK", "nama", "email UNIQUE", "username UNIQUE", "password_hash", "role ENUM(admin)", "is_active", "created_at", "updated_at"]),
    ("produk_emas", "Menyimpan katalog produk emas aktif/nonaktif dan stok.", ["id PK", "kode_produk UNIQUE", "nama_produk", "jenis_emas", "kadar", "berat_gram", "harga_pokok", "stok", "deskripsi", "gambar_url", "status", "created_at", "updated_at", "deleted_at"]),
    ("nasabah", "Menyimpan data nasabah yang dicatat manual oleh admin.", ["id PK", "kode_nasabah UNIQUE", "nama", "no_telepon", "alamat", "catatan", "created_at", "updated_at", "deleted_at"]),
    ("kredit", "Menyimpan transaksi kredit emas beserta snapshot harga, margin, tenor, total terbayar, dan sisa piutang.", ["id PK", "kode_kredit UNIQUE", "nasabah_id FK", "produk_emas_id FK", "tanggal_kredit", "harga_pokok_snapshot", "margin_persen", "margin_nominal", "total_harga_kredit", "tenor_bulan", "periode_angsuran", "jumlah_periode", "nominal_angsuran", "total_terbayar", "sisa_piutang", "status", "catatan", "created_at", "updated_at"]),
    ("jadwal_angsuran", "Menyimpan jadwal tagihan setiap periode untuk satu transaksi kredit.", ["id PK", "kredit_id FK", "angsuran_ke", "tanggal_jatuh_tempo", "nominal_tagihan", "nominal_dibayar", "status", "tanggal_dibayar", "created_at", "updated_at"]),
    ("pembayaran_angsuran", "Menyimpan pembayaran manual yang dicatat admin.", ["id PK", "kode_pembayaran UNIQUE", "kredit_id FK", "jadwal_angsuran_id FK nullable", "tanggal_bayar", "nominal_bayar", "metode_pembayaran", "keterangan", "dicatat_oleh FK users", "created_at", "updated_at"]),
    ("whatsapp_logs", "Menyimpan log pembuatan template dan URL WhatsApp wa.me.", ["id PK", "tipe", "target", "tujuan_nomor", "nama_tujuan", "pesan", "wa_url", "status", "related_type", "related_id", "created_by", "created_at", "updated_at"]),
    ("pengaturan_sistem", "Menyimpan identitas toko, nomor WhatsApp toko, margin default, logo teks, dan alamat.", ["id PK", "nama_toko", "nomor_whatsapp_toko", "margin_default", "logo_text", "alamat_toko", "created_at", "updated_at"]),
]

RELATIONS = [
    "nasabah.id 1..N kredit.nasabah_id",
    "produk_emas.id 1..N kredit.produk_emas_id",
    "kredit.id 1..N jadwal_angsuran.kredit_id",
    "kredit.id 1..N pembayaran_angsuran.kredit_id",
    "jadwal_angsuran.id 1..N pembayaran_angsuran.jadwal_angsuran_id (opsional)",
    "users.id 1..N pembayaran_angsuran.dicatat_oleh",
    "whatsapp_logs.related_type/related_id berelasi logis ke produk_emas, kredit, jadwal_angsuran, atau pembayaran_angsuran",
]

MODULES = [
    ("Publik", "Beranda, katalog produk emas, detail produk, simulasi kredit flat rate, dan pengajuan via WhatsApp tanpa akun pelanggan."),
    ("Autentikasi Admin", "Login, validasi role admin, sesi admin, proteksi filter adminauth, dan logout."),
    ("Produk Emas", "CRUD produk emas, pencarian, status aktif/nonaktif, stok, harga, kadar, dan berat produk."),
    ("Nasabah", "CRUD nasabah, normalisasi nomor WhatsApp, dan kartu piutang nasabah."),
    ("Kredit", "Pembuatan kredit, kalkulasi margin flat rate, generate jadwal angsuran, pengurangan stok, detail kredit, pembatalan, dan template WA transaksi/pengingat/lunas."),
    ("Pembayaran", "Pencatatan pembayaran manual, alokasi ke jadwal angsuran, update total terbayar, sisa piutang, status kredit, dan template WA konfirmasi pembayaran."),
    ("Piutang", "Monitoring kredit aktif, lunas, jatuh tempo hari ini, dan terlambat."),
    ("Laporan", "Laporan kredit, pembayaran, piutang, filter tanggal/status, dan ekspor CSV."),
    ("WhatsApp Logs", "Daftar log template pesan dengan filter tipe dan status."),
    ("Pengaturan", "Pengelolaan nama toko, nomor WhatsApp toko, margin default, logo teks, dan alamat toko."),
]

SOURCES = [
    "README.md",
    "app/Config/Routes.php",
    "app/Filters/AdminAuthFilter.php",
    "app/Helpers/mahen_helper.php",
    "app/Models/*.php",
    "app/Database/Migrations/*.php",
    "app/Database/Seeds/MahenGoldSeeder.php",
    "app/Services/*.php",
    "app/Controllers/PublicController.php",
    "app/Controllers/Admin/*.php",
    "app/Controllers/Api/KreditApiController.php",
    "app/Views/public/*.php",
    "app/Views/admin/**/*.php",
    "app/Views/partials/admin/sidebar.php",
    "public/assets/css/mahengold.css",
    "public/assets/js/mahengold.js",
    "contoh/* sebagai acuan format, bukan sumber isi domain",
]

PAGES_PUBLIC = [
    ("Halaman Beranda Publik", "/", "Landing page MahenGold dengan hero, produk unggulan, keunggulan layanan, dan ajakan membuka katalog atau WhatsApp."),
    ("Halaman Katalog Produk", "/katalog", "Daftar produk emas aktif, pencarian, filter jenis, kartu produk, harga, stok, simulasi bulanan/mingguan, dan tombol detail/pengajuan."),
    ("Halaman Detail Produk", "/produk/MGD-001", "Detail satu produk emas, harga, kadar, berat, stok, deskripsi, simulasi kredit, dan modal pengajuan WhatsApp."),
]

PAGES_ADMIN = [
    ("Halaman Login Admin", "/admin/login", "Form username dan password untuk akun admin."),
    ("Dashboard Admin", "/admin/dashboard", "Kartu statistik nilai kredit, pembayaran, piutang, kredit aktif/lunas, jatuh tempo, terlambat, kredit dan pembayaran terbaru, piutang tertinggi, dan log WhatsApp."),
    ("Manajemen Produk", "/admin/produk", "Daftar produk emas dengan pencarian, tambah, edit, hapus, status, harga, kadar, berat, dan stok."),
    ("Manajemen Nasabah", "/admin/nasabah", "Daftar nasabah, pencarian, tambah, edit, hapus, dan kartu piutang."),
    ("Transaksi Kredit", "/admin/kredit", "Daftar kredit, filter status, detail kredit, dan pembuatan transaksi kredit."),
    ("Detail Kredit", "/admin/kredit/1", "Detail kredit, nasabah, produk, jadwal angsuran, pembayaran, pembatalan, dan tombol WhatsApp."),
    ("Pembayaran Angsuran", "/admin/pembayaran", "Daftar pembayaran manual, kode pembayaran, kredit, nasabah, nominal, metode, dan notifikasi WhatsApp."),
    ("Monitoring Piutang", "/admin/piutang", "Monitoring sisa piutang berdasarkan kredit aktif/lunas/jatuh tempo/terlambat."),
    ("Laporan Kredit", "/admin/laporan/kredit", "Laporan kredit dengan filter dan ekspor CSV."),
    ("Laporan Pembayaran", "/admin/laporan/pembayaran", "Laporan pembayaran angsuran dengan filter tanggal dan ekspor CSV."),
    ("Laporan Piutang", "/admin/laporan/piutang", "Laporan sisa piutang kredit dan statusnya."),
    ("WhatsApp Logs", "/admin/whatsapp-logs", "Log template WhatsApp berdasarkan tipe, tujuan, status, dan URL wa.me."),
    ("Pengaturan Sistem", "/admin/pengaturan", "Form identitas toko, nomor WhatsApp, margin default, logo teks, dan alamat toko."),
]

PUBLIC_SHOT_MAP = {
    "Halaman Beranda Publik": "public_home.png",
    "Halaman Katalog Produk": "public_katalog.png",
    "Halaman Detail Produk": "public_detail_produk.png",
}

ADMIN_SHOT_MAP = {
    "Halaman Login Admin": "admin_login.png",
    "Dashboard Admin": "admin_dashboard.png",
    "Manajemen Produk": "admin_produk.png",
    "Manajemen Nasabah": "admin_nasabah.png",
    "Transaksi Kredit": "admin_kredit.png",
    "Detail Kredit": "admin_kredit_detail.png",
    "Pembayaran Angsuran": "admin_pembayaran.png",
    "Monitoring Piutang": "admin_piutang.png",
    "Laporan Kredit": "admin_laporan_kredit.png",
    "Laporan Pembayaran": "admin_laporan_kredit.png",
    "Laporan Piutang": "admin_laporan_kredit.png",
    "WhatsApp Logs": "admin_whatsapp_logs.png",
    "Pengaturan Sistem": "admin_pengaturan.png",
}

ADMIN_MOCKUP_MAP = {
    "Halaman Login Admin": "admin_login.png",
    "Dashboard Admin": "admin_dashboard.png",
    "Manajemen Produk": "admin_produk.png",
    "Manajemen Nasabah": "admin_nasabah.png",
    "Transaksi Kredit": "admin_kredit.png",
    "Detail Kredit": "admin_kredit_detail.png",
    "Pembayaran Angsuran": "admin_pembayaran.png",
    "Monitoring Piutang": "admin_piutang.png",
    "Laporan Kredit": "admin_laporan_kredit.png",
    "Laporan Pembayaran": "admin_laporan_pembayaran.png",
    "Laporan Piutang": "admin_laporan_piutang.png",
    "WhatsApp Logs": "admin_whatsapp_logs.png",
    "Pengaturan Sistem": "admin_pengaturan.png",
}


def mockup_path(filename: str) -> Path:
    return MOCKUP_DIR / filename.replace(".png", "_mockup_lite.png")


def ensure_dirs() -> None:
    for d in (DIAGRAM_DIR, DOCS_DIR, AUDIT_DIR, SCREENSHOT_DIR, MOCKUP_DIR, PREVIEW_DIR):
        d.mkdir(parents=True, exist_ok=True)


def shape_node(node_id: str, label: str, kind: str, x: int, y: int, w: int = 150, h: int = 70) -> dict[str, Any]:
    return {"id": node_id, "label": label, "kind": kind, "x": x, "y": y, "w": w, "h": h}


def edge(src: str, dst: str, label: str) -> dict[str, str]:
    return {"src": src, "dst": dst, "label": label}


def build_diagrams() -> dict[str, dict[str, Any]]:
    diagrams: dict[str, dict[str, Any]] = {}

    diagrams["Diagram_Konteks"] = {
        "title": "Diagram Konteks Sistem Informasi MahenGold",
        "nodes": [
            shape_node("E1", "Pelanggan / Calon Nasabah", "entity", 60, 170, 190, 70),
            shape_node("P0", "Sistem Informasi\nPenjualan dan Kredit Emas\nMahenGold", "process", 370, 120, 240, 160),
            shape_node("E2", "Admin", "entity", 760, 75, 170, 70),
            shape_node("E3", "WhatsApp wa.me", "entity", 760, 260, 170, 70),
        ],
        "edges": [
            edge("E1", "P0", "data_katalog_simulasi_pengajuan"),
            edge("P0", "E1", "info_produk_simulasi_wa"),
            edge("E2", "P0", "data_login_master_kredit_pembayaran_laporan"),
            edge("P0", "E2", "info_dashboard_master_kredit_piutang_laporan_log"),
            edge("P0", "E3", "data_template_pesan_wa"),
            edge("E3", "P0", "info_url_wa"),
        ],
    }

    nodes_l0 = [
        shape_node("E1", "Pelanggan / Calon Nasabah", "entity", 40, 170, 175, 70),
        shape_node("E2", "Admin", "entity", 40, 440, 175, 70),
        shape_node("E3", "WhatsApp wa.me", "entity", 1060, 320, 175, 70),
        shape_node("P1", "1.0 Layanan Publik\nKatalog & Simulasi", "process", 300, 80, 190, 90),
        shape_node("P2", "2.0 Autentikasi\nAdmin", "process", 300, 220, 190, 90),
        shape_node("P3", "3.0 Pengelolaan\nData Master", "process", 300, 370, 190, 90),
        shape_node("P4", "4.0 Transaksi Kredit\n& Jadwal", "process", 560, 150, 190, 90),
        shape_node("P5", "5.0 Pembayaran\n& Piutang", "process", 560, 320, 190, 90),
        shape_node("P6", "6.0 Laporan &\nKomunikasi WA", "process", 820, 230, 190, 90),
        shape_node("D1", "D1 users", "store", 300, 540, 180, 55),
        shape_node("D2", "D2 produk_emas", "store", 560, 30, 180, 55),
        shape_node("D3", "D3 nasabah", "store", 560, 490, 180, 55),
        shape_node("D4", "D4 kredit", "store", 820, 60, 180, 55),
        shape_node("D5", "D5 jadwal_angsuran", "store", 820, 135, 180, 55),
        shape_node("D6", "D6 pembayaran_angsuran", "store", 820, 410, 205, 55),
        shape_node("D7", "D7 whatsapp_logs", "store", 1060, 430, 180, 55),
        shape_node("D8", "D8 pengaturan_sistem", "store", 560, 590, 205, 55),
    ]
    edges_l0 = [
        edge("E1", "P1", "data_katalog_simulasi_pengajuan"), edge("P1", "E1", "info_produk_simulasi_wa"),
        edge("E2", "P2", "data_login"), edge("P2", "E2", "info_session_admin"),
        edge("E2", "P3", "data_master"), edge("P3", "E2", "info_master"),
        edge("E2", "P4", "data_kredit"), edge("P4", "E2", "info_kredit_jadwal"),
        edge("E2", "P5", "data_pembayaran_piutang"), edge("P5", "E2", "info_pembayaran_piutang"),
        edge("E2", "P6", "data_laporan_wa"), edge("P6", "E2", "info_laporan_log"),
        edge("P6", "E3", "data_template_pesan_wa"), edge("E3", "P6", "info_url_wa"),
        edge("P2", "D1", "data_login_check"), edge("D1", "P2", "info_akun_admin"),
        edge("P1", "D2", "data_produk_query"), edge("D2", "P1", "info_produk"),
        edge("P3", "D2", "data_produk"), edge("D2", "P3", "info_produk"),
        edge("P3", "D3", "data_nasabah"), edge("D3", "P3", "info_nasabah"),
        edge("P4", "D4", "data_kredit"), edge("D4", "P4", "info_kredit"),
        edge("P4", "D5", "data_jadwal"), edge("D5", "P4", "info_jadwal"),
        edge("P5", "D6", "data_pembayaran"), edge("D6", "P5", "info_pembayaran"),
        edge("P6", "D7", "data_log_wa"), edge("D7", "P6", "info_log_wa"),
        edge("P3", "D8", "data_pengaturan"), edge("D8", "P3", "info_pengaturan"),
        edge("P1", "P4", "data_simulasi_kredit"), edge("P4", "P1", "info_estimasi_kredit"),
        edge("P4", "P5", "data_tagihan"), edge("P5", "P6", "info_rekap_pembayaran"),
    ]
    diagrams["DFD_Level0"] = {"title": "DFD Level 0 Sistem Informasi MahenGold", "nodes": nodes_l0, "edges": edges_l0}

    diagrams["DFD_Level1_P1"] = dfd_level1_p1()
    diagrams["DFD_Level1_P2"] = dfd_level1_p2()
    diagrams["DFD_Level1_P3"] = dfd_level1_p3()
    diagrams["DFD_Level1_P4"] = dfd_level1_p4()
    diagrams["DFD_Level1_P5"] = dfd_level1_p5()
    diagrams["DFD_Level1_P6"] = dfd_level1_p6()
    diagrams["ERD_Chen"] = erd_chens()
    diagrams["ERD_CrowsFoot"] = erd_crows()
    diagrams["Konseptual_Basis_Data"] = conceptual_db()
    diagrams["Struktur_Tabel_Relasi"] = table_structure()
    return diagrams


def dfd_level1_p1() -> dict[str, Any]:
    return {"title": "DFD Level 1 Proses 1.0 Layanan Publik Katalog & Simulasi", "nodes": [
        shape_node("E1", "Pelanggan / Calon Nasabah", "entity", 40, 210, 180, 70),
        shape_node("P11", "1.1 Tampilkan\nBeranda", "process", 310, 40, 170, 80),
        shape_node("P12", "1.2 Cari Katalog\nProduk", "process", 310, 160, 170, 80),
        shape_node("P13", "1.3 Detail Produk", "process", 310, 280, 170, 80),
        shape_node("P14", "1.4 Hitung Simulasi\nFlat Rate", "process", 580, 160, 190, 80),
        shape_node("P15", "1.5 Buat Link\nPengajuan WA", "process", 840, 210, 190, 80),
        shape_node("E3", "WhatsApp wa.me", "entity", 1120, 210, 170, 70),
        shape_node("D2", "D2 produk_emas", "store", 575, 40, 185, 55),
        shape_node("D7", "D7 whatsapp_logs", "store", 840, 360, 185, 55),
        shape_node("D8", "D8 pengaturan_sistem", "store", 575, 360, 215, 55),
    ], "edges": [
        edge("E1", "P11", "data_akses_beranda"), edge("P11", "E1", "info_beranda"),
        edge("E1", "P12", "data_filter_produk"), edge("P12", "E1", "info_katalog_produk"),
        edge("E1", "P13", "data_kode_produk"), edge("P13", "E1", "info_detail_produk"),
        edge("E1", "P14", "data_parameter_simulasi"), edge("P14", "E1", "info_simulasi_kredit"),
        edge("E1", "P15", "data_pengajuan_wa"), edge("P15", "E1", "info_url_pengajuan_wa"),
        edge("P15", "E3", "data_template_pengajuan"), edge("E3", "P15", "info_url_wa"),
        edge("P11", "D2", "data_produk_unggulan_query"), edge("D2", "P11", "info_produk_unggulan"),
        edge("P12", "D2", "data_produk_filter"), edge("D2", "P12", "info_produk_filter"),
        edge("P13", "D2", "data_detail_produk_query"), edge("D2", "P13", "info_detail_produk"),
        edge("P11", "D8", "data_pengaturan_query"), edge("D8", "P11", "info_pengaturan"),
        edge("P14", "D8", "data_margin_query"), edge("D8", "P14", "info_margin_default"),
        edge("P15", "D7", "data_log_pengajuan"), edge("D7", "P15", "info_log_pengajuan"),
        edge("P12", "P13", "data_pilihan_produk"), edge("P13", "P14", "data_harga_produk"), edge("P14", "P15", "info_hasil_simulasi"),
    ]}


def dfd_level1_p2() -> dict[str, Any]:
    return {"title": "DFD Level 1 Proses 2.0 Autentikasi Admin", "nodes": [
        shape_node("E2", "Admin", "entity", 60, 150, 160, 70),
        shape_node("P21", "2.1 Tampilkan\nLogin", "process", 330, 70, 170, 80),
        shape_node("P22", "2.2 Validasi\nKredensial", "process", 330, 230, 170, 80),
        shape_node("P23", "2.3 Bentuk / Hapus\nSesi Admin", "process", 610, 150, 190, 80),
        shape_node("D1", "D1 users", "store", 610, 330, 180, 55),
        shape_node("D8", "D8 pengaturan_sistem", "store", 610, 40, 215, 55),
    ], "edges": [
        edge("E2", "P21", "data_akses_login"), edge("P21", "E2", "info_form_login"),
        edge("E2", "P22", "data_login"), edge("P22", "E2", "info_validasi_login"),
        edge("E2", "P23", "data_logout"), edge("P23", "E2", "info_status_sesi"),
        edge("P21", "D8", "data_pengaturan_login"), edge("D8", "P21", "info_pengaturan_login"),
        edge("P22", "D1", "data_kredensial_admin"), edge("D1", "P22", "info_akun_admin"),
        edge("P22", "P23", "data_admin_valid"),
    ]}


def dfd_level1_p3() -> dict[str, Any]:
    return {"title": "DFD Level 1 Proses 3.0 Pengelolaan Data Master", "nodes": [
        shape_node("E2", "Admin", "entity", 50, 200, 160, 70),
        shape_node("P31", "3.1 Kelola\nProduk Emas", "process", 310, 40, 180, 80),
        shape_node("P32", "3.2 Kelola\nNasabah", "process", 310, 190, 180, 80),
        shape_node("P33", "3.3 Kelola\nPengaturan", "process", 310, 340, 180, 80),
        shape_node("P34", "3.4 Lihat Kartu\nPiutang Nasabah", "process", 610, 190, 195, 80),
        shape_node("D2", "D2 produk_emas", "store", 610, 40, 185, 55),
        shape_node("D3", "D3 nasabah", "store", 610, 320, 185, 55),
        shape_node("D8", "D8 pengaturan_sistem", "store", 610, 430, 215, 55),
        shape_node("D4", "D4 kredit", "store", 890, 165, 180, 55),
    ], "edges": [
        edge("E2", "P31", "data_produk"), edge("P31", "E2", "info_produk"),
        edge("E2", "P32", "data_nasabah"), edge("P32", "E2", "info_nasabah"),
        edge("E2", "P33", "data_pengaturan"), edge("P33", "E2", "info_pengaturan"),
        edge("E2", "P34", "data_kartu_piutang"), edge("P34", "E2", "info_kartu_piutang"),
        edge("P31", "D2", "data_produk"), edge("D2", "P31", "info_produk"),
        edge("P32", "D3", "data_nasabah"), edge("D3", "P32", "info_nasabah"),
        edge("P34", "D3", "data_nasabah_piutang"), edge("D3", "P34", "info_nasabah_piutang"),
        edge("P33", "D8", "data_pengaturan"), edge("D8", "P33", "info_pengaturan"),
        edge("P34", "D4", "data_kredit_nasabah"), edge("D4", "P34", "info_kredit_nasabah"),
    ]}


def dfd_level1_p4() -> dict[str, Any]:
    return {"title": "DFD Level 1 Proses 4.0 Transaksi Kredit & Jadwal", "nodes": [
        shape_node("E2", "Admin", "entity", 40, 230, 160, 70),
        shape_node("E3", "WhatsApp wa.me", "entity", 1130, 230, 170, 70),
        shape_node("P41", "4.1 Pilih Nasabah\n& Produk", "process", 300, 70, 190, 80),
        shape_node("P42", "4.2 Preview\nKredit", "process", 300, 220, 190, 80),
        shape_node("P43", "4.3 Simpan Kredit\n& Kurangi Stok", "process", 570, 120, 200, 80),
        shape_node("P44", "4.4 Generate\nJadwal Angsuran", "process", 570, 290, 200, 80),
        shape_node("P45", "4.5 Kirim Info /\nPengingat WA", "process", 850, 220, 200, 80),
        shape_node("D2", "D2 produk_emas", "store", 570, 20, 185, 55),
        shape_node("D3", "D3 nasabah", "store", 570, 420, 185, 55),
        shape_node("D4", "D4 kredit", "store", 850, 70, 180, 55),
        shape_node("D5", "D5 jadwal_angsuran", "store", 850, 370, 205, 55),
        shape_node("D7", "D7 whatsapp_logs", "store", 1080, 410, 185, 55),
    ], "edges": [
        edge("E2", "P41", "data_pilihan_kredit"), edge("P41", "E2", "info_pilihan_kredit"),
        edge("E2", "P42", "data_parameter_kredit"), edge("P42", "E2", "info_preview_kredit"),
        edge("E2", "P43", "data_transaksi_kredit"), edge("P43", "E2", "info_transaksi_kredit"),
        edge("E2", "P45", "data_permintaan_wa_kredit"), edge("P45", "E2", "info_template_wa_kredit"),
        edge("P45", "E3", "data_template_kredit"), edge("E3", "P45", "info_url_wa"),
        edge("P41", "D2", "data_produk_kredit"), edge("D2", "P41", "info_produk_kredit"),
        edge("P43", "D2", "data_stok_produk"), edge("D2", "P43", "info_stok_produk"),
        edge("P41", "D3", "data_nasabah_kredit"), edge("D3", "P41", "info_nasabah_kredit"),
        edge("P43", "D4", "data_kredit"), edge("D4", "P43", "info_kredit"),
        edge("P45", "D4", "data_kredit_wa"), edge("D4", "P45", "info_kredit_wa"),
        edge("P44", "D5", "data_jadwal"), edge("D5", "P44", "info_jadwal"),
        edge("P45", "D5", "data_jadwal_wa"), edge("D5", "P45", "info_jadwal_wa"),
        edge("P45", "D7", "data_log_kredit_wa"), edge("D7", "P45", "info_log_kredit_wa"),
        edge("P41", "P42", "data_produk_nasabah"), edge("P42", "P43", "data_hasil_preview"), edge("P43", "P44", "data_kredit_baru"), edge("P44", "P45", "info_jadwal_baru"),
    ]}


def dfd_level1_p5() -> dict[str, Any]:
    return {"title": "DFD Level 1 Proses 5.0 Pembayaran & Piutang", "nodes": [
        shape_node("E2", "Admin", "entity", 40, 210, 160, 70),
        shape_node("E3", "WhatsApp wa.me", "entity", 1100, 225, 170, 70),
        shape_node("P51", "5.1 Pilih Kredit\nAktif", "process", 300, 60, 190, 80),
        shape_node("P52", "5.2 Catat\nPembayaran", "process", 300, 220, 190, 80),
        shape_node("P53", "5.3 Update Jadwal\n& Status Kredit", "process", 575, 220, 205, 80),
        shape_node("P54", "5.4 Monitoring\nPiutang", "process", 575, 60, 190, 80),
        shape_node("P55", "5.5 Notifikasi\nPembayaran WA", "process", 835, 225, 205, 80),
        shape_node("D4", "D4 kredit", "store", 575, 390, 180, 55),
        shape_node("D5", "D5 jadwal_angsuran", "store", 575, 500, 205, 55),
        shape_node("D6", "D6 pembayaran_angsuran", "store", 835, 390, 225, 55),
        shape_node("D7", "D7 whatsapp_logs", "store", 835, 500, 185, 55),
    ], "edges": [
        edge("E2", "P51", "data_pilih_kredit"), edge("P51", "E2", "info_kredit_aktif"),
        edge("E2", "P52", "data_pembayaran"), edge("P52", "E2", "info_pembayaran"),
        edge("E2", "P54", "data_filter_piutang"), edge("P54", "E2", "info_piutang"),
        edge("E2", "P55", "data_konfirmasi_wa"), edge("P55", "E2", "info_url_konfirmasi_wa"),
        edge("P55", "E3", "data_template_pembayaran"), edge("E3", "P55", "info_url_wa"),
        edge("P51", "D4", "data_kredit_aktif_query"), edge("D4", "P51", "info_kredit_aktif"),
        edge("P53", "D4", "data_status_kredit"), edge("D4", "P53", "info_status_kredit"),
        edge("P54", "D4", "data_piutang_query"), edge("D4", "P54", "info_piutang"),
        edge("P51", "D5", "data_jadwal_query"), edge("D5", "P51", "info_jadwal"),
        edge("P53", "D5", "data_status_jadwal"), edge("D5", "P53", "info_status_jadwal"),
        edge("P52", "D6", "data_pembayaran"), edge("D6", "P52", "info_pembayaran"),
        edge("P55", "D7", "data_log_pembayaran_wa"), edge("D7", "P55", "info_log_pembayaran_wa"),
        edge("P51", "P52", "info_kredit_terpilih"), edge("P52", "P53", "data_alokasi_pembayaran"), edge("P53", "P55", "info_status_pembayaran"),
    ]}


def dfd_level1_p6() -> dict[str, Any]:
    return {"title": "DFD Level 1 Proses 6.0 Laporan & Komunikasi WhatsApp", "nodes": [
        shape_node("E2", "Admin", "entity", 40, 230, 160, 70),
        shape_node("P61", "6.1 Laporan\nKredit", "process", 300, 60, 180, 80),
        shape_node("P62", "6.2 Laporan\nPembayaran", "process", 300, 220, 180, 80),
        shape_node("P63", "6.3 Laporan\nPiutang", "process", 300, 380, 180, 80),
        shape_node("P64", "6.4 Filter\nWhatsApp Logs", "process", 590, 220, 190, 80),
        shape_node("P65", "6.5 Ekspor\nCSV", "process", 850, 220, 170, 80),
        shape_node("D4", "D4 kredit", "store", 590, 45, 180, 55),
        shape_node("D6", "D6 pembayaran_angsuran", "store", 590, 375, 225, 55),
        shape_node("D7", "D7 whatsapp_logs", "store", 850, 375, 185, 55),
    ], "edges": [
        edge("E2", "P61", "data_filter_laporan_kredit"), edge("P61", "E2", "info_laporan_kredit"),
        edge("E2", "P62", "data_filter_laporan_pembayaran"), edge("P62", "E2", "info_laporan_pembayaran"),
        edge("E2", "P63", "data_filter_laporan_piutang"), edge("P63", "E2", "info_laporan_piutang"),
        edge("E2", "P64", "data_filter_log_wa"), edge("P64", "E2", "info_log_wa"),
        edge("E2", "P65", "data_permintaan_csv"), edge("P65", "E2", "info_file_csv"),
        edge("P61", "D4", "data_kredit_laporan"), edge("D4", "P61", "info_kredit_laporan"),
        edge("P63", "D4", "data_piutang_laporan"), edge("D4", "P63", "info_piutang_laporan"),
        edge("P62", "D6", "data_pembayaran_laporan"), edge("D6", "P62", "info_pembayaran_laporan"),
        edge("P64", "D7", "data_log_wa_filter"), edge("D7", "P64", "info_log_wa_filter"),
        edge("P61", "P65", "info_data_kredit_csv"), edge("P62", "P65", "info_data_pembayaran_csv"), edge("P63", "P65", "info_data_piutang_csv"),
    ]}


def erd_chens() -> dict[str, Any]:
    nodes = []
    x_positions = [80, 360, 640, 920]
    y_positions = [70, 220, 370]
    for idx, (name, desc, fields) in enumerate(TABLES):
        x = x_positions[idx % 4]
        y = y_positions[idx // 4]
        nodes.append(shape_node("T" + str(idx + 1), name + "\n" + "\n".join(fields[:7]), "entity", x, y, 225, 120))
    edges_ = [
        edge("T3", "T4", "1:N memiliki_kredit"), edge("T2", "T4", "1:N dikreditkan"), edge("T4", "T5", "1:N memiliki_jadwal"),
        edge("T4", "T6", "1:N memiliki_pembayaran"), edge("T5", "T6", "1:N dibayar"), edge("T1", "T6", "1:N mencatat"),
        edge("T8", "T2", "mengatur_margin_harga"), edge("T7", "T4", "mencatat_pesan"),
    ]
    return {"title": "ERD Chen MahenGold", "nodes": nodes, "edges": edges_}


def erd_crows() -> dict[str, Any]:
    d = erd_chens()
    d["title"] = "ERD Crow's Foot MahenGold"
    for e in d["edges"]:
        e["label"] = e["label"].replace("1:N", "1 --<")
    return d


def conceptual_db() -> dict[str, Any]:
    nodes = [
        shape_node("C1", "Admin\n(users)", "entity", 70, 170, 170, 80),
        shape_node("C2", "Produk Emas", "entity", 330, 70, 170, 80),
        shape_node("C3", "Nasabah", "entity", 330, 270, 170, 80),
        shape_node("C4", "Kredit Emas", "entity", 590, 170, 180, 90),
        shape_node("C5", "Jadwal Angsuran", "entity", 850, 70, 190, 80),
        shape_node("C6", "Pembayaran Angsuran", "entity", 850, 270, 210, 80),
        shape_node("C7", "WhatsApp Logs", "entity", 590, 390, 180, 80),
        shape_node("C8", "Pengaturan Sistem", "entity", 70, 390, 200, 80),
    ]
    edges_ = [edge("C3", "C4", "mengajukan / memiliki"), edge("C2", "C4", "menjadi objek kredit"), edge("C4", "C5", "membentuk"), edge("C4", "C6", "dibayar melalui"), edge("C5", "C6", "dialokasikan ke"), edge("C1", "C6", "mencatat"), edge("C4", "C7", "menjadi referensi pesan"), edge("C8", "C2", "menentukan margin default")]
    return {"title": "Konseptual Basis Data MahenGold", "nodes": nodes, "edges": edges_}


def table_structure() -> dict[str, Any]:
    nodes = []
    for idx, (name, desc, fields) in enumerate(TABLES):
        nodes.append(shape_node("S" + str(idx + 1), name + "\n" + "\n".join(fields), "store", 40 + (idx % 2) * 560, 30 + (idx // 2) * 170, 480, 140))
    edges_ = [edge("S3", "S4", "nasabah_id"), edge("S2", "S4", "produk_emas_id"), edge("S4", "S5", "kredit_id"), edge("S4", "S6", "kredit_id"), edge("S5", "S6", "jadwal_angsuran_id"), edge("S1", "S6", "dicatat_oleh")]
    return {"title": "Struktur Tabel dan Relasi MahenGold", "nodes": nodes, "edges": edges_}


def graphml(diagram: dict[str, Any], path: Path) -> None:
    ns = "http://graphml.graphdrawing.org/xmlns"
    yns = "http://www.yworks.com/xml/graphml"
    ET.register_namespace("", ns)
    ET.register_namespace("y", yns)
    root = ET.Element(f"{{{ns}}}graphml")
    ET.SubElement(root, f"{{{ns}}}key", {"id": "d0", "for": "node", "yfiles.type": "nodegraphics"})
    ET.SubElement(root, f"{{{ns}}}key", {"id": "d1", "for": "edge", "yfiles.type": "edgegraphics"})
    graph = ET.SubElement(root, f"{{{ns}}}graph", {"id": "G", "edgedefault": "directed"})
    for n in diagram["nodes"]:
        node = ET.SubElement(graph, f"{{{ns}}}node", {"id": n["id"]})
        data = ET.SubElement(node, f"{{{ns}}}data", {"key": "d0"})
        sn = ET.SubElement(data, f"{{{yns}}}ShapeNode")
        ET.SubElement(sn, f"{{{yns}}}Geometry", {"x": str(n["x"]), "y": str(n["y"]), "width": str(n["w"]), "height": str(n["h"])})
        fill = {"entity": "#FFF2CC", "process": "#DAE8FC", "store": "#E2F0D9"}.get(n["kind"], "#FFFFFF")
        ET.SubElement(sn, f"{{{yns}}}Fill", {"color": fill, "transparent": "false"})
        ET.SubElement(sn, f"{{{yns}}}BorderStyle", {"color": "#2F3A45", "type": "line", "width": "1.5"})
        label = ET.SubElement(sn, f"{{{yns}}}NodeLabel", {"fontFamily": "Arial", "fontSize": "12", "alignment": "center", "autoSizePolicy": "content"})
        label.text = n["label"]
        shape = "ellipse" if n["kind"] == "process" else "rectangle"
        ET.SubElement(sn, f"{{{yns}}}Shape", {"type": shape})
    for i, e in enumerate(diagram["edges"]):
        ed = ET.SubElement(graph, f"{{{ns}}}edge", {"id": f"e{i}", "source": e["src"], "target": e["dst"]})
        data = ET.SubElement(ed, f"{{{ns}}}data", {"key": "d1"})
        pe = ET.SubElement(data, f"{{{yns}}}PolyLineEdge")
        ET.SubElement(pe, f"{{{yns}}}LineStyle", {"color": "#2F3A45", "type": "line", "width": "1.0"})
        ET.SubElement(pe, f"{{{yns}}}Arrows", {"source": "none", "target": "standard"})
        lab = ET.SubElement(pe, f"{{{yns}}}EdgeLabel", {"fontFamily": "Arial", "fontSize": "10", "configuration": "AutoFlippingLabel"})
        lab.text = e["label"]
    ET.ElementTree(root).write(path, encoding="utf-8", xml_declaration=True)


def drawio(diagram: dict[str, Any], path: Path) -> None:
    mxfile = ET.Element("mxfile", {"host": "app.diagrams.net"})
    diag = ET.SubElement(mxfile, "diagram", {"name": diagram["title"][:60]})
    model = ET.SubElement(diag, "mxGraphModel", {"dx": "1400", "dy": "900", "grid": "1", "gridSize": "10", "guides": "1", "tooltips": "1", "connect": "1", "arrows": "1", "fold": "1", "page": "1", "pageScale": "1", "pageWidth": "1400", "pageHeight": "900", "math": "0", "shadow": "0"})
    root = ET.SubElement(model, "root")
    ET.SubElement(root, "mxCell", {"id": "0"})
    ET.SubElement(root, "mxCell", {"id": "1", "parent": "0"})
    for n in diagram["nodes"]:
        if n["kind"] == "process":
            style = "ellipse;whiteSpace=wrap;html=1;fillColor=#dae8fc;strokeColor=#2f3a45;fontFamily=Arial;fontSize=12;"
        elif n["kind"] == "store":
            style = "shape=datastore;whiteSpace=wrap;html=1;boundedLbl=1;backgroundOutline=1;size=15;fillColor=#e2f0d9;strokeColor=#2f3a45;fontFamily=Arial;fontSize=12;"
        else:
            style = "rounded=0;whiteSpace=wrap;html=1;fillColor=#fff2cc;strokeColor=#2f3a45;fontFamily=Arial;fontSize=12;"
        cell = ET.SubElement(root, "mxCell", {"id": n["id"], "value": html.escape(n["label"]).replace("\n", "<br>"), "style": style, "vertex": "1", "parent": "1"})
        ET.SubElement(cell, "mxGeometry", {"x": str(n["x"]), "y": str(n["y"]), "width": str(n["w"]), "height": str(n["h"]), "as": "geometry"})
    for i, e in enumerate(diagram["edges"]):
        cell = ET.SubElement(root, "mxCell", {"id": f"E{i}", "value": html.escape(e["label"]), "style": "endArrow=block;html=1;rounded=0;strokeColor=#2f3a45;fontFamily=Arial;fontSize=10;", "edge": "1", "parent": "1", "source": e["src"], "target": e["dst"]})
        ET.SubElement(cell, "mxGeometry", {"relative": "1", "as": "geometry"})
    ET.ElementTree(mxfile).write(path, encoding="utf-8", xml_declaration=True)


def preview_png(diagram: dict[str, Any], path: Path) -> None:
    max_x = max(n["x"] + n["w"] for n in diagram["nodes"]) + 80
    max_y = max(n["y"] + n["h"] for n in diagram["nodes"]) + 80
    im = Image.new("RGB", (max(1200, max_x), max(700, max_y)), "white")
    draw = ImageDraw.Draw(im)
    try:
        font = ImageFont.truetype("arial.ttf", 14)
        small = ImageFont.truetype("arial.ttf", 11)
        title_font = ImageFont.truetype("arialbd.ttf", 18)
    except Exception:
        font = small = title_font = ImageFont.load_default()
    draw.text((30, 15), diagram["title"], fill="#111111", font=title_font)
    centers = {n["id"]: (n["x"] + n["w"] / 2, n["y"] + n["h"] / 2) for n in diagram["nodes"]}
    for e in diagram["edges"]:
        x1, y1 = centers[e["src"]]
        x2, y2 = centers[e["dst"]]
        draw.line((x1, y1, x2, y2), fill="#606060", width=1)
        ang = math.atan2(y2 - y1, x2 - x1)
        ax, ay = x2 - math.cos(ang) * 30, y2 - math.sin(ang) * 30
        left = (ax - math.cos(ang - math.pi / 6) * 10, ay - math.sin(ang - math.pi / 6) * 10)
        right = (ax - math.cos(ang + math.pi / 6) * 10, ay - math.sin(ang + math.pi / 6) * 10)
        draw.polygon([(ax, ay), left, right], fill="#606060")
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        draw.rectangle((mx - 4, my - 7, mx + len(e["label"]) * 5 + 4, my + 8), fill="white")
        draw.text((mx, my - 7), e["label"], fill="#444444", font=small)
    for n in diagram["nodes"]:
        x, y, w, h = n["x"], n["y"], n["w"], n["h"]
        fill = {"entity": "#FFF2CC", "process": "#DAE8FC", "store": "#E2F0D9"}.get(n["kind"], "#FFFFFF")
        if n["kind"] == "process":
            draw.ellipse((x, y, x + w, y + h), fill=fill, outline="#2F3A45", width=2)
        else:
            draw.rectangle((x, y, x + w, y + h), fill=fill, outline="#2F3A45", width=2)
            if n["kind"] == "store":
                draw.line((x + 10, y + 12, x + w - 10, y + 12), fill="#2F3A45", width=1)
                draw.line((x + 10, y + h - 12, x + w - 10, y + h - 12), fill="#2F3A45", width=1)
        lines = n["label"].split("\n")
        total_h = len(lines) * 16
        for idx, line in enumerate(lines):
            tw = draw.textlength(line, font=font)
            draw.text((x + (w - tw) / 2, y + (h - total_h) / 2 + idx * 16), line, fill="#111111", font=font)
    im.save(path)


def balance_table(diagram: dict[str, Any]) -> list[tuple[str, list[str], list[str], str]]:
    stores = {n["id"]: n["label"].replace("\n", " ") for n in diagram["nodes"] if n["kind"] == "store"}
    rows = []
    for sid, label in stores.items():
        ins = [e["label"] for e in diagram["edges"] if e["dst"] == sid and e["label"].startswith("data_")]
        outs = [e["label"] for e in diagram["edges"] if e["src"] == sid and e["label"].startswith("info_")]
        rows.append((label, ins, outs, "Seimbang" if len(ins) == len(outs) and len(ins) > 0 else "Tidak"))
    return rows


def validate_xml_files() -> list[str]:
    results = []
    for p in sorted(DIAGRAM_DIR.glob("*.graphml")) + sorted(DIAGRAM_DIR.glob("*.drawio.xml")):
        try:
            ET.parse(p)
            results.append(f"OK XML: {p.relative_to(ROOT)}")
        except Exception as exc:
            results.append(f"ERROR XML: {p.relative_to(ROOT)} -> {exc}")
    return results


def generate_diagrams() -> dict[str, dict[str, Any]]:
    diagrams = build_diagrams()
    for name, d in diagrams.items():
        graphml(d, DIAGRAM_DIR / f"{name}.graphml")
        drawio(d, DIAGRAM_DIR / f"{name}.drawio.xml")
        preview_png(d, PREVIEW_DIR / f"{name}.png")
    return diagrams


def screenshot_fallback(filename: str, title: str, url: str, reason: str = "") -> None:
    im = Image.new("RGB", (1366, 768), "#f8f5ee")
    draw = ImageDraw.Draw(im)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 44)
        font = ImageFont.truetype("arial.ttf", 24)
        small = ImageFont.truetype("arial.ttf", 18)
    except Exception:
        title_font = font = small = ImageFont.load_default()
    draw.rectangle((0, 0, 1366, 90), fill="#111827")
    draw.text((40, 24), "MahenGold", fill="#f7c948", font=title_font)
    draw.rounded_rectangle((70, 150, 1296, 610), radius=24, fill="white", outline="#d4af37", width=3)
    draw.text((110, 200), title, fill="#111827", font=title_font)
    draw.text((110, 285), f"URL aktual: {url}", fill="#374151", font=font)
    draw.text((110, 340), "Fallback visual dibuat karena browser otomatis tidak tersedia.", fill="#6b7280", font=small)
    if reason:
        draw.text((110, 380), reason[:130], fill="#6b7280", font=small)
    draw.text((110, 500), "Halaman ini tetap berasal dari route dan view nyata aplikasi MahenGold.", fill="#374151", font=small)
    im.save(SCREENSHOT_DIR / filename)


def capture_screenshots() -> list[str]:
    public_shots = [
        ("public_home.png", "Halaman Beranda Publik", "/", False),
        ("public_katalog.png", "Halaman Katalog Produk", "/katalog", False),
        ("public_detail_produk.png", "Halaman Detail Produk", "/produk/MGD-001", False),
        ("admin_login.png", "Halaman Login Admin", "/admin/login", False),
    ]
    protected_shots = [
        ("admin_dashboard.png", "Dashboard Admin", "/admin/dashboard", True),
        ("admin_produk.png", "Manajemen Produk", "/admin/produk", True),
        ("admin_nasabah.png", "Manajemen Nasabah", "/admin/nasabah", True),
        ("admin_kredit.png", "Transaksi Kredit", "/admin/kredit", True),
        ("admin_kredit_detail.png", "Detail Kredit", "/admin/kredit/1", True),
        ("admin_pembayaran.png", "Pembayaran Angsuran", "/admin/pembayaran", True),
        ("admin_piutang.png", "Monitoring Piutang", "/admin/piutang", True),
        ("admin_laporan_kredit.png", "Laporan Kredit", "/admin/laporan/kredit", True),
        ("admin_whatsapp_logs.png", "WhatsApp Logs", "/admin/whatsapp-logs", True),
        ("admin_pengaturan.png", "Pengaturan Sistem", "/admin/pengaturan", True),
    ]
    base = "http://localhost:8080"
    captured = []
    try:
        _wait_for_tcp("localhost", 3306, timeout_seconds=20)
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = None
            last_exc: Exception | None = None
            for launch_args in ({"channel": "msedge", "headless": True}, {"headless": True}):
                try:
                    browser = p.chromium.launch(**launch_args)
                    break
                except Exception as exc:
                    last_exc = exc
            if browser is None:
                raise RuntimeError(str(last_exc))
            context = browser.new_context(viewport={"width": 1366, "height": 768}, device_scale_factor=1)
            page = context.new_page()

            for filename, title, route, need_auth in public_shots:
                _goto_valid(page, base + route, title, require_admin=False)
                time.sleep(0.4)
                page.screenshot(path=str(SCREENSHOT_DIR / filename), full_page=False)
                captured.append(f"OK screenshot: screenshots/{filename} -> {route}")

            _goto_valid(page, base + "/admin/login", "Login Admin MahenGold", require_admin=False)
            page.fill('input[name="username"]', "admin")
            page.fill('input[name="password"]', "admin123")
            page.click('button[type="submit"]')
            page.wait_for_load_state("networkidle", timeout=15000)
            if "/admin/dashboard" not in page.url:
                raise RuntimeError(f"Login admin gagal, URL saat ini: {page.url}")

            for filename, title, route, need_auth in protected_shots:
                _goto_valid(page, base + route, title, require_admin=True)
                time.sleep(0.4)
                page.screenshot(path=str(SCREENSHOT_DIR / filename), full_page=False)
                captured.append(f"OK screenshot: screenshots/{filename} -> {route}")
            context.close()
            browser.close()
    except Exception as exc:
        for filename, title, route, _ in public_shots + protected_shots:
            screenshot_fallback(filename, title, base + route, str(exc))
            captured.append(f"Fallback screenshot: screenshots/{filename} -> {route}")
    return captured


def _wait_for_tcp(host: str, port: int, timeout_seconds: int = 20) -> None:
    import socket

    deadline = time.time() + timeout_seconds
    last_error: OSError | None = None
    while time.time() < deadline:
        sock = socket.socket()
        sock.settimeout(2)
        try:
            sock.connect((host, port))
            sock.close()
            return
        except OSError as exc:
            last_error = exc
            sock.close()
            time.sleep(1)
    raise RuntimeError(f"Port {host}:{port} belum dapat diakses untuk koneksi database. Error terakhir: {last_error}")


def _goto_valid(page: Any, url: str, title: str, require_admin: bool = False, retries: int = 3) -> None:
    last_problem = ""
    for attempt in range(1, retries + 1):
        response = page.goto(url, wait_until="networkidle", timeout=20000)
        time.sleep(0.6)
        status = response.status if response is not None else 0
        body_text = page.locator("body").inner_text(timeout=5000)
        current_title = page.title()
        current_url = page.url
        error_markers = [
            "CodeIgniter\\Database\\Exceptions\\DatabaseException",
            "Unable to connect to the database",
            "No connection could be made because the target machine actively refused it",
            "CodeIgniter\\Security\\Exceptions\\SecurityException",
            "The action you requested is not allowed",
        ]
        has_error = status >= 500 or any(marker in body_text for marker in error_markers) or any(marker in current_title for marker in error_markers)
        logged_out = require_admin and "/admin/login" in current_url
        if not has_error and not logged_out:
            return
        last_problem = f"status={status}, title={current_title}, url={current_url}"
        time.sleep(1.2 * attempt)
    raise RuntimeError(f"Gagal membuka halaman {title} tanpa error setelah {retries} percobaan: {last_problem}")


def create_mockup_images(opacity: float = 1.0) -> list[str]:
    """Create 16:9 lite wireframe mockups for BAB 4.4.

    These images are intentionally not implementation screenshots. They use a
    clean lite style, and visual slots are represented with explicit LOGO and
    GAMBAR placeholders as requested.
    """
    canvas_size = (1366, 768)
    specs = [
        ("public_home.png", "Beranda Publik", "public_home"),
        ("public_katalog.png", "Katalog Produk", "public_catalog"),
        ("public_detail_produk.png", "Detail Produk", "public_detail"),
        ("admin_login.png", "Login Admin", "admin_login"),
        ("admin_dashboard.png", "Dashboard Admin", "admin_dashboard"),
        ("admin_produk.png", "Manajemen Produk", "admin_table"),
        ("admin_nasabah.png", "Manajemen Nasabah", "admin_table"),
        ("admin_kredit.png", "Transaksi Kredit", "admin_table"),
        ("admin_kredit_detail.png", "Detail Kredit", "admin_detail"),
        ("admin_pembayaran.png", "Pembayaran Angsuran", "admin_table"),
        ("admin_piutang.png", "Monitoring Piutang", "admin_table"),
        ("admin_laporan_kredit.png", "Laporan", "admin_report"),
        ("admin_laporan_pembayaran.png", "Laporan Pembayaran", "admin_report"),
        ("admin_laporan_piutang.png", "Laporan Piutang", "admin_report"),
        ("admin_whatsapp_logs.png", "WhatsApp Logs", "admin_table"),
        ("admin_pengaturan.png", "Pengaturan Sistem", "admin_form"),
    ]

    logs: list[str] = []
    for filename, title, layout in specs:
        image = Image.new("RGB", canvas_size, "#f8fafc")
        draw = ImageDraw.Draw(image)
        fonts = _mockup_fonts()
        _draw_mockup_background(draw, canvas_size)
        if layout.startswith("public"):
            _draw_public_mockup(draw, layout, title, fonts)
        elif layout == "admin_login":
            _draw_login_mockup(draw, title, fonts)
        else:
            _draw_admin_mockup(draw, layout, title, fonts)
        dst = mockup_path(filename)
        image.save(dst)
        logs.append(f"OK lite mockup 16:9: {dst.relative_to(ROOT)}")
    return logs


def _mockup_fonts() -> dict[str, ImageFont.FreeTypeFont | ImageFont.ImageFont]:
    try:
        return {
            "title": ImageFont.truetype("arialbd.ttf", 34),
            "h2": ImageFont.truetype("arialbd.ttf", 24),
            "body": ImageFont.truetype("arial.ttf", 18),
            "small": ImageFont.truetype("arial.ttf", 14),
            "placeholder": ImageFont.truetype("arialbd.ttf", 28),
        }
    except Exception:
        font = ImageFont.load_default()
        return {"title": font, "h2": font, "body": font, "small": font, "placeholder": font}


def _draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font: ImageFont.ImageFont, fill: str = "#334155") -> None:
    left, top, right, bottom = box
    bbox = draw.textbbox((0, 0), text, font=font)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.text((left + ((right - left) - width) / 2, top + ((bottom - top) - height) / 2), text, fill=fill, font=font)


def _draw_mockup_background(draw: ImageDraw.ImageDraw, size: tuple[int, int]) -> None:
    width, height = size
    for x in range(0, width, 32):
        draw.line((x, 0, x, height), fill="#eef2f7", width=1)
    for y in range(0, height, 32):
        draw.line((0, y, width, y), fill="#eef2f7", width=1)
    draw.rounded_rectangle((34, 28, width - 34, height - 28), radius=28, fill="#ffffff", outline="#dbe3ee", width=2)


def _placeholder(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fonts: dict[str, ImageFont.ImageFont]) -> None:
    draw.rounded_rectangle(box, radius=18, fill="#f1f5f9", outline="#94a3b8", width=2)
    left, top, right, bottom = box
    draw.line((left + 18, top + 18, right - 18, bottom - 18), fill="#cbd5e1", width=2)
    draw.line((right - 18, top + 18, left + 18, bottom - 18), fill="#cbd5e1", width=2)
    _draw_centered_text(draw, box, text, fonts["placeholder"], "#64748b")


def _button(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fonts: dict[str, ImageFont.ImageFont], fill: str = "#f8d57e") -> None:
    draw.rounded_rectangle(box, radius=16, fill=fill, outline="#d4a72c", width=2)
    _draw_centered_text(draw, box, text, fonts["small"], "#111827")


def _input(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], label: str, fonts: dict[str, ImageFont.ImageFont]) -> None:
    draw.rounded_rectangle(box, radius=12, fill="#ffffff", outline="#cbd5e1", width=2)
    draw.text((box[0] + 18, box[1] + 13), label, fill="#94a3b8", font=fonts["small"])


def _draw_nav(draw: ImageDraw.ImageDraw, fonts: dict[str, ImageFont.ImageFont]) -> None:
    draw.rounded_rectangle((64, 54, 1302, 118), radius=18, fill="#ffffff", outline="#dbe3ee", width=2)
    _placeholder(draw, (86, 68, 168, 104), "LOGO", fonts)
    draw.text((190, 74), "MahenGold", fill="#0f172a", font=fonts["h2"])
    for idx, item in enumerate(["Beranda", "Katalog", "Simulasi", "WhatsApp"]):
        draw.text((730 + idx * 120, 78), item, fill="#475569", font=fonts["small"])


def _draw_public_mockup(draw: ImageDraw.ImageDraw, layout: str, title: str, fonts: dict[str, ImageFont.ImageFont]) -> None:
    _draw_nav(draw, fonts)
    draw.text((86, 145), title, fill="#0f172a", font=fonts["title"])
    if layout == "public_home":
        draw.text((88, 210), "Hero penjualan dan kredit emas", fill="#475569", font=fonts["body"])
        _button(draw, (88, 270, 260, 322), "Lihat Katalog", fonts)
        _button(draw, (280, 270, 452, 322), "Ajukan WA", fonts, "#e2e8f0")
        _placeholder(draw, (760, 165, 1240, 430), "GAMBAR", fonts)
        for i, label in enumerate(["Produk Unggulan", "Simulasi Kredit", "WhatsApp"]):
            x = 88 + i * 390
            draw.rounded_rectangle((x, 490, x + 340, 650), radius=20, fill="#ffffff", outline="#dbe3ee", width=2)
            _placeholder(draw, (x + 24, 512, x + 120, 594), "GAMBAR", fonts)
            draw.text((x + 140, 530), label, fill="#334155", font=fonts["body"])
            draw.line((x + 140, 565, x + 300, 565), fill="#cbd5e1", width=5)
            draw.line((x + 140, 595, x + 260, 595), fill="#e2e8f0", width=5)
    elif layout == "public_catalog":
        _input(draw, (86, 165, 520, 220), "Cari produk emas", fonts)
        _input(draw, (540, 165, 790, 220), "Filter jenis", fonts)
        _button(draw, (810, 165, 930, 220), "Filter", fonts)
        for row in range(2):
            for col in range(3):
                x = 86 + col * 405
                y = 260 + row * 220
                draw.rounded_rectangle((x, y, x + 350, y + 180), radius=20, fill="#ffffff", outline="#dbe3ee", width=2)
                _placeholder(draw, (x + 24, y + 24, x + 148, y + 130), "GAMBAR", fonts)
                draw.text((x + 168, y + 30), "Nama Produk", fill="#334155", font=fonts["body"])
                draw.line((x + 168, y + 70, x + 310, y + 70), fill="#cbd5e1", width=5)
                draw.line((x + 168, y + 105, x + 285, y + 105), fill="#e2e8f0", width=5)
                _button(draw, (x + 168, y + 130, x + 310, y + 164), "Detail", fonts)
    else:
        _placeholder(draw, (86, 170, 570, 575), "GAMBAR", fonts)
        draw.text((630, 170), "Nama Produk Emas", fill="#0f172a", font=fonts["title"])
        for y in [240, 280, 320, 360]:
            draw.line((635, y, 1135, y), fill="#cbd5e1", width=7)
        draw.rounded_rectangle((630, 420, 1180, 595), radius=18, fill="#f8fafc", outline="#dbe3ee", width=2)
        draw.text((660, 445), "Panel Simulasi Kredit", fill="#334155", font=fonts["h2"])
        _input(draw, (660, 490, 870, 542), "Tenor", fonts)
        _input(draw, (890, 490, 1110, 542), "Periode", fonts)
        _button(draw, (660, 560, 860, 612), "Ajukan WA", fonts)


def _draw_admin_shell(draw: ImageDraw.ImageDraw, title: str, fonts: dict[str, ImageFont.ImageFont]) -> None:
    draw.rounded_rectangle((64, 54, 285, 714), radius=22, fill="#111827", outline="#111827", width=2)
    _placeholder(draw, (92, 82, 160, 132), "LOGO", fonts)
    draw.text((92, 155), "Admin Console", fill="#e5e7eb", font=fonts["small"])
    for idx, item in enumerate(["Dashboard", "Produk", "Nasabah", "Kredit", "Pembayaran", "Piutang", "Laporan", "WA Logs"]):
        y = 205 + idx * 48
        fill = "#374151" if idx == 0 else "#1f2937"
        draw.rounded_rectangle((92, y, 256, y + 34), radius=10, fill=fill)
        draw.text((112, y + 8), item, fill="#f8fafc", font=fonts["small"])
    draw.rounded_rectangle((315, 54, 1302, 120), radius=18, fill="#ffffff", outline="#dbe3ee", width=2)
    draw.text((345, 72), title, fill="#0f172a", font=fonts["h2"])
    draw.rounded_rectangle((1160, 72, 1260, 104), radius=12, fill="#f1f5f9", outline="#cbd5e1", width=1)
    draw.text((1180, 80), "Admin", fill="#475569", font=fonts["small"])


def _draw_admin_mockup(draw: ImageDraw.ImageDraw, layout: str, title: str, fonts: dict[str, ImageFont.ImageFont]) -> None:
    _draw_admin_shell(draw, title, fonts)
    if layout == "admin_dashboard":
        for i, label in enumerate(["Kredit", "Pembayaran", "Piutang", "Jatuh Tempo"]):
            x = 335 + i * 235
            draw.rounded_rectangle((x, 155, x + 205, 250), radius=18, fill="#ffffff", outline="#dbe3ee", width=2)
            draw.text((x + 22, 178), label, fill="#475569", font=fonts["small"])
            draw.line((x + 22, 215, x + 160, 215), fill="#f8d57e", width=9)
        _table(draw, (335, 300, 815, 655), "Kredit Terbaru", fonts)
        _table(draw, (850, 300, 1260, 655), "Jadwal / Log", fonts)
    elif layout == "admin_detail":
        draw.rounded_rectangle((335, 155, 760, 330), radius=18, fill="#ffffff", outline="#dbe3ee", width=2)
        draw.text((365, 180), "Informasi Kredit", fill="#334155", font=fonts["h2"])
        for y in [230, 265, 300]:
            draw.line((365, y, 720, y), fill="#cbd5e1", width=6)
        draw.rounded_rectangle((790, 155, 1260, 330), radius=18, fill="#ffffff", outline="#dbe3ee", width=2)
        draw.text((820, 180), "Aksi WhatsApp", fill="#334155", font=fonts["h2"])
        _button(draw, (820, 235, 980, 285), "Kirim Info", fonts)
        _button(draw, (1000, 235, 1180, 285), "Pengingat", fonts, "#e2e8f0")
        _table(draw, (335, 370, 1260, 655), "Jadwal Angsuran dan Pembayaran", fonts)
    elif layout == "admin_report":
        _input(draw, (335, 155, 560, 208), "Tanggal mulai", fonts)
        _input(draw, (585, 155, 810, 208), "Tanggal akhir", fonts)
        _button(draw, (835, 155, 970, 208), "Filter", fonts)
        _button(draw, (990, 155, 1140, 208), "Export CSV", fonts, "#e2e8f0")
        _table(draw, (335, 250, 1260, 655), title, fonts)
    elif layout == "admin_form":
        draw.rounded_rectangle((335, 155, 900, 650), radius=18, fill="#ffffff", outline="#dbe3ee", width=2)
        draw.text((365, 180), "Form Pengaturan", fill="#334155", font=fonts["h2"])
        for idx, label in enumerate(["Nama toko", "Nomor WhatsApp", "Margin default", "Logo text", "Alamat toko"]):
            _input(draw, (365, 235 + idx * 70, 840, 285 + idx * 70), label, fonts)
        _button(draw, (365, 590, 560, 642), "Simpan", fonts)
    else:
        _input(draw, (335, 155, 760, 208), "Cari / Filter data", fonts)
        _button(draw, (790, 155, 960, 208), "Tambah", fonts)
        _table(draw, (335, 250, 1260, 655), title, fonts)


def _draw_login_mockup(draw: ImageDraw.ImageDraw, title: str, fonts: dict[str, ImageFont.ImageFont]) -> None:
    _placeholder(draw, (590, 105, 776, 180), "LOGO", fonts)
    draw.text((575, 205), title, fill="#0f172a", font=fonts["title"])
    draw.rounded_rectangle((438, 275, 928, 585), radius=28, fill="#ffffff", outline="#dbe3ee", width=2)
    _input(draw, (500, 335, 866, 390), "Username", fonts)
    _input(draw, (500, 415, 866, 470), "Password", fonts)
    _button(draw, (500, 505, 866, 560), "Masuk Admin", fonts)


def _table(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, fonts: dict[str, ImageFont.ImageFont]) -> None:
    left, top, right, bottom = box
    draw.rounded_rectangle(box, radius=18, fill="#ffffff", outline="#dbe3ee", width=2)
    draw.text((left + 24, top + 20), title, fill="#334155", font=fonts["h2"])
    header_y = top + 72
    draw.rounded_rectangle((left + 24, header_y, right - 24, header_y + 42), radius=10, fill="#f1f5f9")
    for col in range(4):
        x = left + 48 + col * ((right - left - 96) // 4)
        draw.line((x, header_y + 18, x + 110, header_y + 18), fill="#94a3b8", width=5)
    row_y = header_y + 62
    while row_y < bottom - 30:
        draw.line((left + 24, row_y, right - 24, row_y), fill="#e2e8f0", width=2)
        for col in range(4):
            x = left + 48 + col * ((right - left - 96) // 4)
            draw.line((x, row_y + 22, x + 130, row_y + 22), fill="#cbd5e1", width=5)
        row_y += 48


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)


def set_run(run, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def p(doc: Document, text: str = "", align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent: bool = True) -> Any:
    par = doc.add_paragraph()
    par.alignment = align
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    if indent:
        par.paragraph_format.first_line_indent = Inches(0.35)
    run = par.add_run(text)
    set_run(run)
    return par


def heading(doc: Document, text: str, level: int = 1) -> None:
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    run = par.add_run(text)
    set_run(run, bold=True, size=11 if level else 12)


def center_heading(doc: Document, text: str, size: int = 12) -> None:
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    run = par.add_run(text)
    set_run(run, bold=True, size=size)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        run = table.rows[0].cells[i].paragraphs[0].add_run(h)
        set_run(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val))
            set_run(run)


def caption(doc: Document, text: str) -> None:
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    run = par.add_run(text)
    set_run(run, italic=True)


def add_image(doc: Document, path: Path, caption_text: str, width: float = 6.2) -> None:
    if not path.exists():
        return
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    par.add_run().add_picture(str(path), width=Inches(width))
    caption(doc, caption_text)


def start_doc() -> Document:
    doc = Document()
    configure_doc(doc)
    center_heading(doc, "BAB IV", 12)
    center_heading(doc, "HASIL DAN PEMBAHASAN", 12)
    return doc


def doc_41() -> Document:
    doc = start_doc()
    heading(doc, "4.1 Analisis")
    p(doc, f"Sistem yang dianalisis adalah {SYSTEM['domain']}. Berdasarkan pembacaan source code aktual, sistem dibangun sebagai aplikasi web back-office untuk penjualan produk emas dan pencatatan kredit emas secara manual oleh admin. Area publik tidak menyediakan akun pelanggan; calon nasabah hanya melihat katalog, menghitung simulasi kredit flat rate, lalu meneruskan pengajuan melalui tautan WhatsApp wa.me.")
    heading(doc, "4.1.1 Analisis Kebutuhan Fungsional")
    p(doc, "Kebutuhan fungsional disusun berdasarkan role yang benar-benar terdapat pada kode program. Role yang tersimpan di tabel users hanya admin, sedangkan pengunjung publik diperlakukan sebagai pelanggan atau calon nasabah tanpa proses login.")
    add_table(doc, ["Role", "Kebutuhan Fungsional"], [
        ["Admin", "Login, mengelola produk emas, nasabah, transaksi kredit, jadwal angsuran, pembayaran, piutang, laporan, log WhatsApp, dan pengaturan sistem."],
        ["Pelanggan/Calon Nasabah", "Melihat beranda, katalog, detail produk, simulasi kredit, serta membuat pengajuan awal melalui WhatsApp tanpa akun pelanggan."],
    ])
    heading(doc, "4.1.2 Analisis Kebutuhan Nonfungsional")
    add_table(doc, ["Aspek", "Kebutuhan"], [
        ["Platform", "Aplikasi web berbasis CodeIgniter 4, PHP 8, Bootstrap 5, JavaScript, dan MariaDB/MySQL."],
        ["Keamanan", "Area admin diproteksi filter adminauth dan password diverifikasi dengan password_verify."],
        ["Usability", "Antarmuka menggunakan layout admin, sidebar, kartu statistik, tabel, pencarian, filter, dan tombol aksi."],
        ["Integrasi", "WhatsApp menggunakan URL wa.me dan log template pesan; tidak memakai payment gateway atau upload bukti bayar."],
        ["Data", "Data utama disimpan dalam delapan tabel aplikasi dan relasi foreign key pada kredit, jadwal, pembayaran, produk, nasabah, dan users."],
    ])
    heading(doc, "4.1.3 Perangkat Keras dan Perangkat Lunak")
    add_table(doc, ["No", "Perangkat", "Keterangan"], [["1", "Komputer/Laptop", "Digunakan untuk menjalankan editor, browser, PHP, dan database lokal."], ["2", "Sistem Operasi", "Windows 11 sesuai lingkungan kerja saat audit."], ["3", "Web Server", "PHP spark serve atau Apache/XAMPP."], ["4", "Database", "MariaDB/MySQL."], ["5", "Editor", "Visual Studio Code."]])
    heading(doc, "4.1.4 Analisis Proses Bisnis")
    p(doc, "Alur bisnis dimulai ketika calon nasabah membuka beranda dan katalog produk emas, memilih produk, lalu menggunakan fitur simulasi kredit. Hasil simulasi dapat dikirim sebagai template pengajuan WhatsApp kepada admin. Admin kemudian melakukan login, mencatat nasabah, mengelola produk, membuat transaksi kredit, dan sistem menghitung margin flat rate serta membentuk jadwal angsuran. Setelah pembayaran dilakukan di luar sistem, admin mencatat pembayaran manual, sistem memperbarui jadwal angsuran, total terbayar, sisa piutang, dan status kredit. Admin dapat memantau piutang, membuat notifikasi WhatsApp, melihat log pesan, serta mengekspor laporan kredit, pembayaran, dan piutang.")
    return doc


def doc_42(diagrams: dict[str, dict[str, Any]]) -> Document:
    doc = start_doc()
    heading(doc, "4.2 Desain Sistem")
    p(doc, "Desain sistem disusun dalam bentuk Diagram Konteks, DFD Level 0, dan DFD Level 1 untuk setiap proses induk. Seluruh arus data distandardisasi dengan awalan data_ untuk input dan info_ untuk output. Data store dibuat seimbang secara count, yaitu jumlah arus masuk data_ sama dengan jumlah arus keluar info_ pada setiap data store.")
    fig = 1
    for name in ["Diagram_Konteks", "DFD_Level0", "DFD_Level1_P1", "DFD_Level1_P2", "DFD_Level1_P3", "DFD_Level1_P4", "DFD_Level1_P5", "DFD_Level1_P6"]:
        heading(doc, {"Diagram_Konteks": "4.2.1 Diagram Konteks", "DFD_Level0": "4.2.2 DFD Level 0"}.get(name, f"4.2.{fig+2} {diagrams[name]['title']}"))
        p(doc, f"{diagrams[name]['title']} menggambarkan aliran data yang berasal dari fungsi nyata pada aplikasi MahenGold. Entitas, proses, dan data store pada diagram mengikuti route, controller, service, model, migration, dan view yang tersedia pada source code.")
        add_image(doc, PREVIEW_DIR / f"{name}.png", f"Gambar 4.{fig} {diagrams[name]['title']}")
        fig += 1
    heading(doc, "4.2.9 Audit Balancing Data Store")
    for name in ["DFD_Level0", "DFD_Level1_P1", "DFD_Level1_P2", "DFD_Level1_P3", "DFD_Level1_P4", "DFD_Level1_P5", "DFD_Level1_P6"]:
        heading(doc, name.replace("_", " "), 2)
        rows = []
        for label, ins, outs, status in balance_table(diagrams[name]):
            rows.append([label, ", ".join(ins) + f" (×{len(ins)})", ", ".join(outs) + f" (×{len(outs)})", status])
        add_table(doc, ["Database", "Input (data_)", "Output (info_)", "Balance"], rows)
    return doc


def doc_43() -> Document:
    doc = start_doc()
    heading(doc, "4.3 Perancangan Basis Data")
    p(doc, "Perancangan basis data MahenGold didasarkan pada file migration, model, query controller, service transaksi, dan seeder. Basis data memuat data admin, produk emas, nasabah, kredit, jadwal angsuran, pembayaran angsuran, log WhatsApp, serta pengaturan sistem.")
    heading(doc, "4.3.1 Konseptual Basis Data")
    add_image(doc, PREVIEW_DIR / "Konseptual_Basis_Data.png", "Gambar 4.1 Konseptual Basis Data MahenGold")
    heading(doc, "4.3.2 Struktur Tabel")
    for idx, (name, desc, fields) in enumerate(TABLES, start=1):
        heading(doc, f"4.3.2.{idx} Tabel {name}", 2)
        p(doc, desc)
        add_table(doc, ["Field"], [[f] for f in fields])
    heading(doc, "4.3.3 Relasi Antar Entitas")
    add_table(doc, ["No", "Relasi"], [[str(i), r] for i, r in enumerate(RELATIONS, start=1)])
    heading(doc, "4.3.4 ERD")
    add_image(doc, PREVIEW_DIR / "ERD_Chen.png", "Gambar 4.2 ERD Chen MahenGold")
    add_image(doc, PREVIEW_DIR / "ERD_CrowsFoot.png", "Gambar 4.3 ERD Crow's Foot MahenGold")
    add_image(doc, PREVIEW_DIR / "Struktur_Tabel_Relasi.png", "Gambar 4.4 Struktur Tabel dan Relasi")
    return doc


def doc_445() -> Document:
    doc = start_doc()
    heading(doc, "4.4 Perancangan Antarmuka")
    p(doc, "Perancangan antarmuka disusun berdasarkan halaman yang benar-benar tersedia pada folder view dan route aplikasi MahenGold. Setiap rancangan pada subbab 4.4 dilengkapi mockup versi lite berukuran 16:9. Elemen visual pada rancangan sengaja disederhanakan dengan placeholder bertuliskan LOGO untuk area logo dan GAMBAR untuk area gambar agar berbeda dari screenshot implementasi pada subbab 4.5. Sistem tidak memiliki registrasi/login pelanggan, keranjang, checkout online, payment gateway, atau upload bukti pembayaran sehingga halaman tersebut tidak dimasukkan.")
    heading(doc, "4.4.1 Halaman Publik")
    mockup_fig = 1
    for letter, (title, route, desc) in zip("abc", PAGES_PUBLIC):
        heading(doc, f"{letter}. {title}", 2)
        p(doc, f"{title} dirancang untuk route {route}. {desc} Tata letak menggunakan identitas visual MahenGold, navigasi publik, kartu produk, dan komponen aksi menuju detail atau WhatsApp sesuai kebutuhan halaman.")
        add_image(doc, mockup_path(PUBLIC_SHOT_MAP[title]), f"Gambar 4.4.{mockup_fig} Mockup Lite {title}")
        mockup_fig += 1
    heading(doc, "4.4.2 Halaman Admin")
    for idx, (title, route, desc) in enumerate(PAGES_ADMIN):
        heading(doc, f"{chr(ord('a') + idx)}. {title}", 2)
        p(doc, f"{title} dirancang untuk route {route}. {desc} Area admin menggunakan layout sidebar, header, tabel data, filter, tombol aksi, dan form sesuai modul yang diakses.")
        add_image(doc, mockup_path(ADMIN_MOCKUP_MAP[title]), f"Gambar 4.4.{mockup_fig} Mockup Lite {title}")
        mockup_fig += 1
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    heading(doc, "4.5 Implementasi Sistem")
    p(doc, "Implementasi sistem didokumentasikan melalui screenshot halaman nyata yang diambil dari aplikasi berjalan pada server lokal. Screenshot disusun mengikuti urutan halaman publik dan halaman admin yang tersedia pada source code.")
    heading(doc, "4.5.1 Implementasi Halaman Publik")
    fig = 1
    for title, route, desc in PAGES_PUBLIC:
        heading(doc, title, 2)
        p(doc, f"Pada implementasi, {title.lower()} menampilkan {desc.lower()} Pengguna dapat mengakses halaman ini melalui route {route} tanpa login.")
        add_image(doc, SCREENSHOT_DIR / PUBLIC_SHOT_MAP[title], f"Gambar 4.5.{fig} Implementasi {title}")
        fig += 1
    heading(doc, "4.5.2 Implementasi Halaman Admin")
    for title, route, desc in PAGES_ADMIN:
        heading(doc, title, 2)
        p(doc, f"Pada implementasi, {title.lower()} digunakan oleh admin untuk {desc.lower()} Halaman ini diproteksi autentikasi admin kecuali halaman login.")
        add_image(doc, SCREENSHOT_DIR / ADMIN_SHOT_MAP[title], f"Gambar 4.5.{fig} Implementasi {title}")
        fig += 1
    return doc


def save_doc(doc: Document, path: Path) -> None:
    doc.save(path)


def combine_docs() -> Document:
    doc = start_doc()
    heading(doc, "4.1 Analisis")
    p(doc, f"MahenGold merupakan {SYSTEM['domain']}. Role nyata yang tersedia adalah admin dan pelanggan/calon nasabah tanpa login. Sistem mendukung katalog produk emas, simulasi kredit, pengajuan WhatsApp, pengelolaan produk, nasabah, kredit, pembayaran, piutang, laporan, log WhatsApp, dan pengaturan.")
    heading(doc, "4.2 Desain")
    p(doc, "Desain sistem terdiri dari Diagram Konteks, DFD Level 0 dengan enam proses utama, dan enam DFD Level 1. Seluruh data store telah dibuat balanced dengan jumlah input data_ dan output info_ yang sama.")
    add_image(doc, PREVIEW_DIR / "Diagram_Konteks.png", "Gambar 4.1 Diagram Konteks")
    add_image(doc, PREVIEW_DIR / "DFD_Level0.png", "Gambar 4.2 DFD Level 0")
    heading(doc, "4.3 Perancangan Basis Data")
    p(doc, "Basis data terdiri atas delapan tabel utama: users, produk_emas, nasabah, kredit, jadwal_angsuran, pembayaran_angsuran, whatsapp_logs, dan pengaturan_sistem.")
    add_image(doc, PREVIEW_DIR / "ERD_CrowsFoot.png", "Gambar 4.3 ERD Crow's Foot")
    heading(doc, "4.4 Perancangan Antarmuka")
    p(doc, "Antarmuka dirancang berdasarkan halaman publik dan halaman admin yang tersedia pada view aplikasi MahenGold. Halaman yang tidak ada pada program, seperti registrasi pelanggan, keranjang, checkout, dan upload bukti pembayaran, tidak dimasukkan.")
    heading(doc, "4.5 Implementasi Sistem")
    p(doc, "Implementasi sistem terdokumentasi melalui screenshot halaman nyata dari aplikasi berjalan, meliputi halaman publik, login admin, dashboard, modul master, transaksi kredit, pembayaran, piutang, laporan, WhatsApp logs, dan pengaturan.")
    add_image(doc, SCREENSHOT_DIR / "public_home.png", "Gambar 4.4 Implementasi Beranda Publik")
    add_image(doc, SCREENSHOT_DIR / "admin_dashboard.png", "Gambar 4.5 Implementasi Dashboard Admin")
    return doc


def generate_docs(diagrams: dict[str, dict[str, Any]]) -> None:
    save_doc(doc_41(), DOCS_DIR / "BAB_4_1_Analisis.docx")
    save_doc(doc_42(diagrams), DOCS_DIR / "BAB_4_2_Desain.docx")
    save_doc(doc_43(), DOCS_DIR / "BAB_4_3_Perancangan_Basis_Data.docx")
    save_doc(doc_445(), DOCS_DIR / "4.4.4.5.docx")
    save_doc(combine_docs(), DOCS_DIR / "BAB_4_1_sampai_4_5.docx")


def write_audits(diagrams: dict[str, dict[str, Any]], screenshot_log: list[str]) -> None:
    (AUDIT_DIR / "AUDIT_PROGRAM.md").write_text("\n".join([
        "# AUDIT PROGRAM MAHENGOLD",
        "",
        f"## Domain Sistem\n{SYSTEM['domain']}.",
        "",
        "## Role Nyata\n- admin (tersimpan pada users.role dengan ENUM admin)\n- pelanggan/calon_nasabah tanpa login, hanya menggunakan halaman publik dan WhatsApp",
        "",
        "## Modul Nyata",
        *[f"- **{m}**: {d}" for m, d in MODULES],
        "",
        "## Tabel Nyata",
        *[f"- **{name}**: {desc}" for name, desc, fields in TABLES],
        "",
        "## Relasi Utama",
        *[f"- {r}" for r in RELATIONS],
        "",
        "## Alur Bisnis Utama\nCalon nasabah melihat katalog, menjalankan simulasi kredit, dan mengirim pengajuan WhatsApp. Admin login, mengelola produk dan nasabah, membuat kredit, sistem membentuk jadwal angsuran dan mengurangi stok. Pembayaran dicatat manual oleh admin, jadwal dan sisa piutang diperbarui, lalu admin dapat memantau piutang, membuat notifikasi WhatsApp, dan mengekspor laporan.",
        "",
        "## Sumber Audit",
        *[f"- {s}" for s in SOURCES],
    ]), encoding="utf-8")
    lines = ["# AUDIT DFD BALANCING", "", "## Struktur Level 0", SYSTEM["level0_reason"], ""]
    lines += ["## Parent-Child Level 1", "- P1 memiliki DFD_Level1_P1", "- P2 memiliki DFD_Level1_P2", "- P3 memiliki DFD_Level1_P3", "- P4 memiliki DFD_Level1_P4", "- P5 memiliki DFD_Level1_P5", "- P6 memiliki DFD_Level1_P6", ""]
    for name in ["DFD_Level0", "DFD_Level1_P1", "DFD_Level1_P2", "DFD_Level1_P3", "DFD_Level1_P4", "DFD_Level1_P5", "DFD_Level1_P6"]:
        lines += [f"## {name}", "| Database | Input (data_) | Output (info_) | Balance |", "|---|---|---|---|"]
        for label, ins, outs, status in balance_table(diagrams[name]):
            lines.append(f"| {label} | {', '.join(ins)} (×{len(ins)}) | {', '.join(outs)} (×{len(outs)}) | {status} |")
        lines.append("")
    lines += ["## Konsistensi", "- Diagram konteks selaras dengan DFD Level 0 melalui entitas pelanggan/calon nasabah, admin, dan WhatsApp wa.me.", "- Setiap proses induk Level 0 memiliki child Level 1.", "- Semua label arus eksternal dan data store menggunakan awalan data_ atau info_.", "- Tidak ada data store yang hanya menerima input atau hanya mengeluarkan output."]
    (AUDIT_DIR / "AUDIT_DFD_BALANCING.md").write_text("\n".join(lines), encoding="utf-8")
    (AUDIT_DIR / "AUDIT_DOC_BAB4.md").write_text("\n".join([
        "# AUDIT DOKUMEN BAB 4",
        "",
        "## Sumber Isi Subbab",
        "- 4.1 bersumber dari README, routes, controllers, filters, helpers, services, models, migrations, dan views.",
        "- 4.2 bersumber dari struktur proses aktual pada route/controller/service serta audit balancing data store.",
        "- 4.3 bersumber utama dari migration dan model.",
        "- 4.4 bersumber dari halaman view yang benar-benar tersedia.",
        "- 4.5 bersumber dari screenshot implementasi halaman berjalan.",
        "",
        "## Halaman yang Dimasukkan",
        *[f"- {title} ({route})" for title, route, _ in PAGES_PUBLIC + PAGES_ADMIN],
        "",
        "## Halaman yang Tidak Dimasukkan",
        "- Registrasi pelanggan, login pelanggan, keranjang, checkout online, upload bukti pembayaran, verifikasi pembayaran otomatis, dan payment gateway tidak dimasukkan karena tidak ada pada source code MahenGold.",
        "",
        "## Screenshot",
        *[f"- {x}" for x in screenshot_log],
    ]), encoding="utf-8")
    manifest = ["# OUTPUT MANIFEST", ""]
    for folder in (DIAGRAM_DIR, DOCS_DIR, SCREENSHOT_DIR, AUDIT_DIR):
        for pth in sorted(folder.rglob("*")):
            if pth.is_file():
                manifest.append(f"- `{pth.relative_to(ROOT).as_posix()}` — file output dokumentasi MahenGold")
    (AUDIT_DIR / "OUTPUT_MANIFEST.md").write_text("\n".join(manifest), encoding="utf-8")


def validate_outputs(diagrams: dict[str, dict[str, Any]]) -> None:
    lines = ["# VALIDATION RESULTS", ""]
    lines += validate_xml_files()
    lines.append("")
    for name in ["DFD_Level0", "DFD_Level1_P1", "DFD_Level1_P2", "DFD_Level1_P3", "DFD_Level1_P4", "DFD_Level1_P5", "DFD_Level1_P6"]:
        bad = [row for row in balance_table(diagrams[name]) if row[3] != "Seimbang"]
        lines.append(f"{name}: {'OK balanced' if not bad else 'NOT balanced'}")
    for pth in sorted(DOCS_DIR.glob("*.docx")):
        try:
            with zipfile.ZipFile(pth) as zf:
                zf.testzip()
            lines.append(f"OK DOCX: {pth.relative_to(ROOT)}")
        except Exception as exc:
            lines.append(f"ERROR DOCX: {pth.relative_to(ROOT)} -> {exc}")
    (AUDIT_DIR / "VALIDATION_RESULTS.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    ensure_dirs()
    diagrams = generate_diagrams()
    screenshot_log = capture_screenshots()
    mockup_log = create_mockup_images()
    screenshot_log.extend(mockup_log)
    generate_docs(diagrams)
    write_audits(diagrams, screenshot_log)
    validate_outputs(diagrams)
    print("Generated BAB 4 documentation package for MahenGold")


if __name__ == "__main__":
    main()
